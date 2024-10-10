from datetime import datetime
import os
import uuid
from docx import Document
import asyncio
import logging
 
logger=logging.getLogger("GenLiteApp")
 
 
def get_time_stamp1():
    current_time_stamp = datetime.now()
    datestamp = current_time_stamp.strftime("%d%b%Y")
    timestamp = current_time_stamp.strftime("%H%M")
    unique_id=str(uuid.uuid4())
    return datestamp+"_"+timestamp+"__"+unique_id
 
async def download_doc(form):
    file_dir="/home/LogFiles"
    ts=get_time_stamp1()
    file_path =  os.path.join(file_dir,'FSD_GenLite_Template.docx')
    doc = Document(file_path)

    file_name = f"GenLite_FSD_{ts}.docx"
    download_data=form.ecosystem_context.data
    save_file_path =  os.path.join(file_dir,file_name)
    doc.save(save_file_path)
    return save_file_path
 
async def create_fdd_doc(form):
    ui_data = form.ui_functional_design.data
    service_data = form.services_functional_design.data
    design_data = form.data_functional_design.data
 
    cwd=os.getcwd()
    file_dir= os.path.join(cwd, "downloadtemplates")
    file_path =  os.path.join(file_dir,'FSD_GenLite_Template.docx')
    doc = Document(file_path)
    for paragraph in doc.paragraphs:
        if '[FD_UI_COMPONENT]' in paragraph.text:
            paragraph.text = ui_data.replace("\r","")
        if '[FD_SERVICE_COMPONENT]' in paragraph.text:
            paragraph.text = service_data.replace("\r","")
        if '[FD_DATA_COMPONENT]' in paragraph.text:
            paragraph.text = design_data.replace("\r","")
   
    #file_dir= cwd
    file_dir="/home/LogFiles"
    ts=get_time_stamp1()
    file_name = f"GenLite_FSD_{ts}.docx"
    save_file_path =  os.path.join(file_dir,file_name)
    doc.save(save_file_path)
    return save_file_path

async def create_hld_doc(form):
    ui_data = form.ui_high_level_design.data
    service_data = form.services_high_level_design.data
    design_data = form.data_high_level_design.data
 
    cwd=os.getcwd()
    file_dir= os.path.join(cwd, "downloadtemplates")
    file_path =  os.path.join(file_dir,'HLDD_GenLite_Template.docx')
    doc = Document(file_path)
    for paragraph in doc.paragraphs:
        if '[HLD_UI_COMPONENT]' in paragraph.text:
            paragraph.text = ui_data.replace("\r","")
        if '[HLD_SERVICE_COMPONENT]' in paragraph.text:
            paragraph.text = service_data.replace("\r","")
        if '[HLD_DATA_COMPONENT]' in paragraph.text:
            paragraph.text = design_data.replace("\r","")
   
    #file_dir= cwd
    file_dir="/home/LogFiles"
    ts=get_time_stamp1()
    file_name = f"GenLite_HLDD_{ts}.docx"
    save_file_path =  os.path.join(file_dir,file_name)
    doc.save(save_file_path)
    return save_file_path

async def create_dld_doc(form):
    ui_data = form.ui_detailed_design.data
    service_data = form.services_detailed_design.data
    design_data = form.data_detailed_design.data
 
    cwd=os.getcwd()
    file_dir= os.path.join(cwd, "downloadtemplates")
    file_path =  os.path.join(file_dir,'DDD_GenLite_Template.docx')
    doc = Document(file_path)
    for paragraph in doc.paragraphs:
        if '[DDD_UI_COMPONENT]' in paragraph.text:
            paragraph.text = ui_data.replace("\r","")
        if '[DDD_SERVICE_COMPONENT]' in paragraph.text:
            paragraph.text = service_data.replace("\r","")
        if '[DDD_DATA_COMPONENT]' in paragraph.text:
            paragraph.text = design_data.replace("\r","")
   
    #file_dir= cwd
    file_dir="/home/LogFiles"
    ts=get_time_stamp1()
    file_name = f"GenLite_DDD_{ts}.docx"
    save_file_path =  os.path.join(file_dir,file_name)
    doc.save(save_file_path)
    return save_file_path